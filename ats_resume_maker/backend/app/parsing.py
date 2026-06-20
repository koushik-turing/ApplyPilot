"""Turn an uploaded file (PDF / DOCX / TXT) into plain text."""
from __future__ import annotations
import io
import os


def extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        return _pdf(data)
    if ext == ".docx":
        return _docx(data)
    if ext in (".txt", ".text", ".md", ""):
        return data.decode("utf-8", "ignore").strip()
    if ext == ".doc":
        raise ValueError("Legacy .doc isn't supported — please save as PDF, DOCX, or TXT.")
    raise ValueError(f"Unsupported file type '{ext}'. Upload a PDF, DOCX, or TXT file.")


def _pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ValueError("PDF support needs the 'pdfplumber' package.")
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    text = "\n".join(pages).strip()
    if not text:
        raise ValueError(
            "No text found in the PDF — it may be a scanned image. "
            "Please upload a text-based PDF or a DOCX/TXT file."
        )
    return text


def _docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ValueError("DOCX support needs the 'python-docx' package.")
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(c.text for c in row.cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("The DOCX file appears to be empty.")
    return text
