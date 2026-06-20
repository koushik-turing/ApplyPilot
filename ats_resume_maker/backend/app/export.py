"""Render a structured resume to an industry-standard, ATS-safe document.

ATS-safe by design: single column, standard system fonts, real selectable text
(no images/icons/text-boxes), standard section headings, simple bullets, no tables
for layout. Polished by design: clean header, ruled section headings, right-aligned
dates, consistent spacing — the look recruiters expect, that parsers still read.
"""
from __future__ import annotations
import io

from .schemas import Resume

SECTION_ORDER = ("summary", "skills", "experience", "education", "certifications", "projects")
HEADINGS = {
    "summary": "Professional Summary",
    "skills": "Skills",
    "experience": "Experience",
    "education": "Education",
    "certifications": "Certifications",
    "projects": "Projects",
}

# Selectable, ATS-safe templates. Only colors/typography differ — every template stays
# single-column, real-text, no tables/graphics, so parse-ability is identical.
#   accent  = name + heading color (RGB)   |   rule = heading underline color
TEMPLATES = {
    "classic": {"accent": (31, 41, 59),   "rule": (176, 183, 195), "name": (17, 24, 39),
                "name_size": 20, "heading_rule": True,  "name_align": "L"},
    "modern":  {"accent": (37, 99, 235),  "rule": (37, 99, 235),   "name": (17, 24, 39),
                "name_size": 22, "heading_rule": True,  "name_align": "L"},
    "executive": {"accent": (15, 23, 42), "rule": (148, 163, 184), "name": (15, 23, 42),
                "name_size": 24, "heading_rule": True,  "name_align": "C"},
}
DEFAULT_TEMPLATE = "modern"


def _style(template: str | None) -> dict:
    return TEMPLATES.get((template or DEFAULT_TEMPLATE).lower(), TEMPLATES[DEFAULT_TEMPLATE])


def _contact(r: Resume) -> str:
    return "  |  ".join(x for x in [r.email, r.phone, r.location, *r.links] if x)


# ---------------- plain text ----------------
def to_text(r: Resume) -> str:
    out: list[str] = []
    if r.name:
        out.append(r.name)
    if _contact(r):
        out.append(_contact(r))
    if r.summary:
        out += ["", HEADINGS["summary"].upper(), r.summary]
    if r.skills:
        out += ["", HEADINGS["skills"].upper(), ", ".join(r.skills)]
    if r.experience:
        out += ["", HEADINGS["experience"].upper()]
        for e in r.experience:
            head = " | ".join(x for x in [e.title, e.company] if x)
            dates = " - ".join(x for x in [e.start_date, e.end_date] if x)
            out.append(f"{head}{('  (' + dates + ')') if dates else ''}".strip())
            out += [f"- {b}" for b in e.bullets]
            out.append("")
    if r.education:
        out += [HEADINGS["education"].upper()]
        for ed in r.education:
            out.append(" | ".join(x for x in [ed.degree, ed.institution, ed.year] if x))
    if r.certifications:
        out += ["", HEADINGS["certifications"].upper()] + [f"- {c}" for c in r.certifications]
    if r.projects:
        out += ["", HEADINGS["projects"].upper()] + [f"- {p}" for p in r.projects]
    return "\n".join(out).strip() + "\n"


# ---------------- DOCX ----------------
def to_docx(r: Resume, template: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    st = _style(template)
    accent = RGBColor(*st["accent"])
    rule_hex = "%02X%02X%02X" % st["rule"]
    name_rgb = RGBColor(*st["name"])

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(2)
    content_width = Inches(7.0)  # 8.5" - 2*0.8" margins, used for right-aligned tab

    def _tight(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        return p

    def heading(text: str):
        p = _tight(doc.add_paragraph())
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = accent
        # full-width bottom border = the heading rule (template-colored)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "2"), ("w:color", rule_hex)):
            bottom.set(qn(k), v)
        pbdr.append(bottom)
        pPr.append(pbdr)

    center = (st["name_align"] == "C")

    # ---- header ----
    if r.name:
        p = _tight(doc.add_paragraph())
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(r.name)
        run.bold = True
        run.font.size = Pt(st["name_size"])
        run.font.color.rgb = name_rgb
    if _contact(r):
        p = _tight(doc.add_paragraph())
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(_contact(r))
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x44, 0x4A, 0x57)

    if r.summary:
        heading(HEADINGS["summary"])
        _tight(doc.add_paragraph(r.summary))

    if r.skills:
        heading(HEADINGS["skills"])
        _tight(doc.add_paragraph(" •  ".join(r.skills)))

    if r.experience:
        heading(HEADINGS["experience"])
        for e in r.experience:
            p = _tight(doc.add_paragraph())
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
            title = " — ".join(x for x in [e.title, e.company] if x)
            p.add_run(title).bold = True
            dates = " – ".join(x for x in [e.start_date, e.end_date] if x)
            if dates:
                run = p.add_run("\t" + dates)
                run.italic = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x55, 0x5B, 0x68)
            for b in e.bullets:
                _tight(doc.add_paragraph(b, style="List Bullet"))

    if r.education:
        heading(HEADINGS["education"])
        for ed in r.education:
            p = _tight(doc.add_paragraph())
            p.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
            left = " — ".join(x for x in [ed.degree, ed.institution] if x)
            p.add_run(left).bold = bool(ed.degree)
            if ed.year:
                run = p.add_run("\t" + ed.year)
                run.italic = True
                run.font.size = Pt(9.5)

    if r.certifications:
        heading(HEADINGS["certifications"])
        for c in r.certifications:
            _tight(doc.add_paragraph(c, style="List Bullet"))

    if r.projects:
        heading(HEADINGS["projects"])
        for pj in r.projects:
            _tight(doc.add_paragraph(pj, style="List Bullet"))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------- PDF ----------------
def to_pdf(r: Resume, template: str | None = None) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    st = _style(template)
    accent, rule, name_rgb = st["accent"], st["rule"], st["name"]

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.set_margins(15, 14, 15)
    pdf.add_page()

    def txt(s: str) -> str:
        return (s.replace("—", "-").replace("–", "-").replace("’", "'")
                 .replace("“", '"').replace("”", '"').replace("•", "-").replace("·", "-")
                 .encode("latin-1", "replace").decode("latin-1"))

    def cell(text: str, h: float = 5.0, align: str = "L"):
        pdf.multi_cell(0, h, txt(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align=align)

    def hr(gap: float = 1.0):
        y = pdf.get_y() + gap
        pdf.set_draw_color(*rule)
        pdf.set_line_width(0.4)
        pdf.line(pdf.l_margin, y, pdf.l_margin + pdf.epw, y)
        pdf.ln(gap + 1.6)

    def heading(t: str):
        pdf.ln(2.5)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*accent)
        pdf.cell(0, 5.5, txt(t.upper()), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if st["heading_rule"]:
            hr()
        else:
            pdf.ln(1.0)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)

    def role_line(title: str, dates: str):
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(17, 24, 39)
        if dates:
            pdf.cell(pdf.epw * 0.70, 5.5, txt(title))
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(85, 91, 104)
            pdf.cell(0, 5.5, txt(dates), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.cell(0, 5.5, txt(title), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)

    def bullet(text: str):
        pdf.set_x(pdf.l_margin + 3)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, txt("-  " + text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ---- header ----
    if r.name:
        pdf.set_font("Helvetica", "B", st["name_size"])
        pdf.set_text_color(*name_rgb)
        cell(r.name, h=9.0, align=st["name_align"])
        pdf.set_text_color(0, 0, 0)
    if _contact(r):
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(68, 74, 87)
        cell(_contact(r), h=5, align=st["name_align"])
        pdf.set_text_color(0, 0, 0)
    # accent rule under the header for a polished, modern finish
    pdf.ln(0.5)
    pdf.set_draw_color(*accent)
    pdf.set_line_width(0.6)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + pdf.epw, pdf.get_y())
    pdf.ln(1.0)

    if r.summary:
        heading(HEADINGS["summary"])
        cell(r.summary)
    if r.skills:
        heading(HEADINGS["skills"])
        cell(" -  ".join(r.skills))
    if r.experience:
        heading(HEADINGS["experience"])
        for e in r.experience:
            title = " - ".join(x for x in [e.title, e.company] if x)
            dates = " - ".join(x for x in [e.start_date, e.end_date] if x)
            role_line(title, dates)
            for b in e.bullets:
                bullet(b)
            pdf.ln(1.5)
    if r.education:
        heading(HEADINGS["education"])
        for ed in r.education:
            left = " - ".join(x for x in [ed.degree, ed.institution] if x)
            role_line(left, ed.year)
    if r.certifications:
        heading(HEADINGS["certifications"])
        for c in r.certifications:
            bullet(c)
    if r.projects:
        heading(HEADINGS["projects"])
        for pj in r.projects:
            bullet(pj)

    return bytes(pdf.output())
